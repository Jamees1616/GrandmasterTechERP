from pathlib import Path

app = Path("app.py")
requirements = Path("requirements.txt")

# ============================================================
# BACKUP CURRENT FILES
# ============================================================

app.with_name(
    "app.py.before_full_reports_upgrade"
).write_text(
    app.read_text()
)

# ============================================================
# ADD EXPORT LIBRARIES TO REQUIREMENTS
# ============================================================

req = requirements.read_text()

for package in [
    "openpyxl==3.1.5",
    "reportlab==5.0.0",
    "python-docx==1.2.0",
]:
    package_name = package.split("==")[0]

    if package_name not in req:
        req += "\n" + package

requirements.write_text(
    req.strip() + "\n"
)

# ============================================================
# LOCATE EXISTING REPORT ROUTES
# ============================================================

text = app.read_text()

start = text.find("@app.route('/reports')")
end = text.find("@app.route('/knowledge')")

if start == -1 or end == -1:
    raise SystemExit(
        "ERROR: Could not safely locate report routes."
    )

print("Part 1 complete.")
print("Backup created.")
print("Export libraries checked.")
print("Report routes located.")

# ============================================================
# REPORT ARGUMENT HELPERS
# ============================================================

new_routes = r'''
# ============================================================
# REPORT DOWNLOAD / EXPORT ENGINE
# Supports Monthly, Quarterly, Yearly, All-Time
# Formats: CSV, Excel, PDF, Word
# Currency: Malawian Kwacha (MWK / MK)
# ============================================================

def _report_args():
    from datetime import datetime

    now = datetime.utcnow()

    report_type = request.args.get(
        'report_type',
        'monthly'
    ).lower()

    year = request.args.get(
        'year',
        now.year,
        type=int
    )

    month = request.args.get(
        'month',
        now.month,
        type=int
    )

    quarter = request.args.get(
        'quarter',
        1,
        type=int
    )

    if report_type not in [
        'monthly',
        'quarterly',
        'yearly',
        'all'
    ]:
        report_type = 'monthly'

    if month < 1 or month > 12:
        month = now.month

    if quarter < 1 or quarter > 4:
        quarter = 1

    return (
        report_type,
        year,
        month,
        quarter
    )


def _report_filename(
    report_type,
    year,
    month,
    quarter,
    extension
):
    period = get_report_date_range(
        report_type=report_type,
        year=year,
        month=month,
        quarter=quarter
    )

    label = (
        period['label']
        .replace(' ', '_')
        .replace('/', '-')
    )

    return (
        f"Grandmaster_Tech_ERP_Report_"
        f"{label}.{extension}"
    )


def _get_selected_report():
    report_type, year, month, quarter = _report_args()

    report = get_report_data(
        report_type=report_type,
        year=year,
        month=month,
        quarter=quarter
    )

    return (
        report_type,
        year,
        month,
        quarter,
        report
    )

'''

print("Part 2 complete.")
print("Report helper functions added.")

# ============================================================
# REPORTS PAGE
# ============================================================

new_routes += r'''
@app.route('/reports')
def reports():
    (
        report_type,
        selected_year,
        selected_month,
        selected_quarter,
        report
    ) = _get_selected_report()

    return render_template(
        'reports.html',

        report_type=report_type,
        selected_month=selected_month,
        selected_year=selected_year,
        selected_quarter=selected_quarter,

        report_period_label=report['period']['label'],

        total_jobs=report['total_jobs'],
        total_customers=report['total_customers'],
        total_equipment=report['total_equipment'],

        pending_jobs=report['pending_jobs'],
        in_progress_jobs=report['in_progress_jobs'],
        completed_jobs=report['completed_jobs'],
        completion_rate=report['completion_rate'],

        total_service_charges=report['total_service_charges'],
        total_parts_cost=report['total_parts_cost'],
        total_revenue=report['total_revenue'],
        total_amount_paid=report['total_amount_paid'],
        total_balance=report['total_balance'],
        total_profit=report['total_profit'],

        priority_counts=report['priority_counts'],
        category_counts=report['category_counts'],
        top_customers=report['top_customers'],

        gender_counts=report['gender_counts'],
        age_group_counts=report['age_group_counts'],
        district_counts=report['district_counts'],
        customer_type_counts=report['customer_type_counts']
    )


# ============================================================
# CSV EXPORT
# ============================================================

@app.route('/reports/download/csv')
def download_reports_csv():
    import csv
    from io import StringIO
    from flask import Response

    (
        report_type,
        year,
        month,
        quarter,
        report
    ) = _get_selected_report()

    output = StringIO()
    writer = csv.writer(output)

    period = report['period']

    writer.writerow([
        'GRANDMASTER TECH ERP'
    ])

    writer.writerow([
        'Service & Repair Performance Report'
    ])

    writer.writerow([
        'Report Period',
        period['label']
    ])

    writer.writerow([
        'Report Type',
        period['type'].title()
    ])

    writer.writerow([])

    writer.writerow([
        'SUMMARY'
    ])

    writer.writerow([
        'Total Jobs',
        report['total_jobs']
    ])

    writer.writerow([
        'Completed Jobs',
        report['completed_jobs']
    ])

    writer.writerow([
        'Pending Jobs',
        report['pending_jobs']
    ])

    writer.writerow([
        'In Progress Jobs',
        report['in_progress_jobs']
    ])

    writer.writerow([
        'Completion Rate',
        f"{report['completion_rate']}%"
    ])

    writer.writerow([
        'Customers',
        report['total_customers']
    ])

    writer.writerow([
        'Equipment',
        report['total_equipment']
    ])

    writer.writerow([])

    writer.writerow([
        'FINANCIAL PERFORMANCE - MWK'
    ])

    writer.writerow([
        'Service Charges',
        report['total_service_charges']
    ])

    writer.writerow([
        'Parts Cost',
        report['total_parts_cost']
    ])

    writer.writerow([
        'Total Revenue',
        report['total_revenue']
    ])

    writer.writerow([
        'Amount Paid',
        report['total_amount_paid']
    ])

    writer.writerow([
        'Outstanding Balance',
        report['total_balance']
    ])

    writer.writerow([
        'Profit',
        report['total_profit']
    ])

    writer.writerow([])

    writer.writerow([
        'CUSTOMER DEMOGRAPHICS'
    ])

    writer.writerow([
        'Gender',
        'Count'
    ])

    for key, value in report[
        'gender_counts'
    ].items():
        writer.writerow([
            key,
            value
        ])

    writer.writerow([])

    writer.writerow([
        'Age Group',
        'Count'
    ])

    for key, value in report[
        'age_group_counts'
    ].items():
        writer.writerow([
            key,
            value
        ])

    writer.writerow([])

    writer.writerow([
        'District',
        'Count'
    ])

    for key, value in report[
        'district_counts'
    ].items():
        writer.writerow([
            key,
            value
        ])

    writer.writerow([])

    writer.writerow([
        'Customer Type',
        'Count'
    ])

    for key, value in report[
        'customer_type_counts'
    ].items():
        writer.writerow([
            key,
            value
        ])

    writer.writerow([])

    writer.writerow([
        'JOB DETAILS'
    ])

    writer.writerow([
        'ID',
        'Job Title',
        'Equipment Category',
        'Brand',
        'Model',
        'Status',
        'Priority',
        'Created Date',
        'Service Charge (MWK)',
        'Parts Cost (MWK)',
        'Total Cost (MWK)',
        'Amount Paid (MWK)',
        'Balance (MWK)',
        'Profit (MWK)'
    ])

    for job in report['jobs']:

        equipment = job.equipment

        writer.writerow([
            job.id,
            job.title,
            equipment.category
            if equipment else '',
            equipment.brand
            if equipment else '',
            equipment.model
            if equipment else '',
            job.status,
            job.priority,
            job.created_at.strftime(
                '%Y-%m-%d'
            )
            if job.created_at
            else '',
            job.service_charge or 0,
            job.parts_cost or 0,
            job.total_cost or 0,
            job.amount_paid or 0,
            job.balance or 0,
            job.profit or 0
        ])

    filename = _report_filename(
        report_type,
        year,
        month,
        quarter,
        'csv'
    )

    return Response(
        output.getvalue(),
        mimetype='text/csv',
        headers={
            'Content-Disposition':
            f'attachment; filename="{filename}"'
        }
    )

'''

print("Part 3 complete.")
print("Reports route added.")
print("CSV export upgraded for all report periods.")

# ============================================================
# EXCEL EXPORT
# ============================================================

new_routes += r'''
@app.route('/reports/download/excel')
def download_reports_excel():
    from io import BytesIO
    from flask import send_file
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment

    (
        report_type,
        year,
        month,
        quarter,
        report
    ) = _get_selected_report()

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Report"

    period = report['period']

    rows = [
        ['GRANDMASTER TECH ERP'],
        ['Service & Repair Performance Report'],
        ['Report Period', period['label']],
        ['Report Type', period['type'].title()],
        [],
        ['SUMMARY'],
        ['Total Jobs', report['total_jobs']],
        ['Completed Jobs', report['completed_jobs']],
        ['Pending Jobs', report['pending_jobs']],
        ['In Progress Jobs', report['in_progress_jobs']],
        ['Completion Rate', f"{report['completion_rate']}%"],
        ['Customers', report['total_customers']],
        ['Equipment', report['total_equipment']],
        [],
        ['FINANCIAL PERFORMANCE - MWK'],
        ['Service Charges', report['total_service_charges']],
        ['Parts Cost', report['total_parts_cost']],
        ['Total Revenue', report['total_revenue']],
        ['Amount Paid', report['total_amount_paid']],
        ['Outstanding Balance', report['total_balance']],
        ['Profit', report['total_profit']],
        [],
        ['CUSTOMER DEMOGRAPHICS'],
        ['Gender', 'Count']
    ]

    for key, value in report['gender_counts'].items():
        rows.append([key, value])

    rows += [
        [],
        ['Age Group', 'Count']
    ]

    for key, value in report['age_group_counts'].items():
        rows.append([key, value])

    rows += [
        [],
        ['District', 'Count']
    ]

    for key, value in report['district_counts'].items():
        rows.append([key, value])

    rows += [
        [],
        ['Customer Type', 'Count']
    ]

    for key, value in report['customer_type_counts'].items():
        rows.append([key, value])

    rows += [
        [],
        ['JOB DETAILS'],
        [
            'ID',
            'Job Title',
            'Equipment Category',
            'Brand',
            'Model',
            'Status',
            'Priority',
            'Created Date',
            'Service Charge (MWK)',
            'Parts Cost (MWK)',
            'Total Cost (MWK)',
            'Amount Paid (MWK)',
            'Balance (MWK)',
            'Profit (MWK)'
        ]
    ]

    for job in report['jobs']:

        equipment = job.equipment

        rows.append([
            job.id,
            job.title,
            equipment.category if equipment else '',
            equipment.brand if equipment else '',
            equipment.model if equipment else '',
            job.status,
            job.priority,
            job.created_at.strftime('%Y-%m-%d')
            if job.created_at else '',
            float(job.service_charge or 0),
            float(job.parts_cost or 0),
            float(job.total_cost or 0),
            float(job.amount_paid or 0),
            float(job.balance or 0),
            float(job.profit or 0)
        ])

    for row in rows:
        sheet.append(row)

    for cell in sheet[1]:
        cell.font = Font(
            bold=True,
            size=16
        )

    for cell in sheet[2]:
        cell.font = Font(
            bold=True
        )

    for row in sheet.iter_rows():
        for cell in row:
            cell.alignment = Alignment(
                vertical='top'
            )

    for column in sheet.columns:
        max_length = 0
        column_letter = column[0].column_letter

        for cell in column:
            if cell.value is not None:
                max_length = max(
                    max_length,
                    len(str(cell.value))
                )

        sheet.column_dimensions[
            column_letter
        ].width = min(
            max(max_length + 2, 12),
            40
        )

    output = BytesIO()

    workbook.save(output)

    output.seek(0)

    filename = _report_filename(
        report_type,
        year,
        month,
        quarter,
        'xlsx'
    )

    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype=(
            'application/vnd.openxmlformats-officedocument.'
            'spreadsheetml.sheet'
        )
    )


# ============================================================
# PDF EXPORT
# ============================================================

@app.route('/reports/download/pdf')
def download_reports_pdf():
    from io import BytesIO
    from flask import send_file
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle
    )
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet

    (
        report_type,
        year,
        month,
        quarter,
        report
    ) = _get_selected_report()

    output = BytesIO()

    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=36,
        leftMargin=36,
        topMargin=36,
        bottomMargin=36
    )

    styles = getSampleStyleSheet()

    story = []

    period = report['period']

    story.append(
        Paragraph(
            'GRANDMASTER TECH ERP',
            styles['Title']
        )
    )

    story.append(
        Paragraph(
            'Service & Repair Performance Report',
            styles['Heading2']
        )
    )

    story.append(
        Paragraph(
            f"Reporting Period: {period['label']}",
            styles['Normal']
        )
    )

    story.append(Spacer(1, 18))

    summary_data = [
        ['Metric', 'Value'],
        ['Total Jobs', report['total_jobs']],
        ['Completed Jobs', report['completed_jobs']],
        ['Pending Jobs', report['pending_jobs']],
        ['In Progress Jobs', report['in_progress_jobs']],
        ['Completion Rate', f"{report['completion_rate']}%"],
        ['Customers', report['total_customers']],
        ['Equipment', report['total_equipment']],
    ]

    table = Table(
        summary_data,
        colWidths=[250, 150]
    )

    table.setStyle(
        TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('PADDING', (0, 0), (-1, -1), 6),
        ])
    )

    story.append(table)

    story.append(Spacer(1, 18))

    story.append(
        Paragraph(
            'Financial Performance (MWK)',
            styles['Heading2']
        )
    )

    financial_data = [
        ['Metric', 'Amount (MWK)'],
        ['Service Charges', f"{report['total_service_charges']:,.2f}"],
        ['Parts Cost', f"{report['total_parts_cost']:,.2f}"],
        ['Total Revenue', f"{report['total_revenue']:,.2f}"],
        ['Amount Paid', f"{report['total_amount_paid']:,.2f}"],
        ['Outstanding Balance', f"{report['total_balance']:,.2f}"],
        ['Profit', f"{report['total_profit']:,.2f}"],
    ]

    financial_table = Table(
        financial_data,
        colWidths=[250, 150]
    )

    financial_table.setStyle(
        TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('PADDING', (0, 0), (-1, -1), 6),
        ])
    )

    story.append(financial_table)

    story.append(Spacer(1, 18))

    story.append(
        Paragraph(
            'Customer Demographics',
            styles['Heading2']
        )
    )

    demographic_data = [
        ['Category', 'Value', 'Count']
    ]

    for key, value in report[
        'gender_counts'
    ].items():
        demographic_data.append([
            'Gender',
            key,
            value
        ])

    for key, value in report[
        'age_group_counts'
    ].items():
        demographic_data.append([
            'Age Group',
            key,
            value
        ])

    for key, value in report[
        'district_counts'
    ].items():
        demographic_data.append([
            'District',
            key,
            value
        ])

    for key, value in report[
        'customer_type_counts'
    ].items():
        demographic_data.append([
            'Customer Type',
            key,
            value
        ])

    demographic_table = Table(
        demographic_data,
        colWidths=[120, 180, 100]
    )

    demographic_table.setStyle(
        TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('PADDING', (0, 0), (-1, -1), 5),
        ])
    )

    story.append(demographic_table)

    document.build(story)

    output.seek(0)

    filename = _report_filename(
        report_type,
        year,
        month,
        quarter,
        'pdf'
    )

    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype='application/pdf'
    )


# ============================================================
# WORD EXPORT
# ============================================================

@app.route('/reports/download/word')
def download_reports_word():
    from io import BytesIO
    from flask import send_file
    from docx import Document

    (
        report_type,
        year,
        month,
        quarter,
        report
    ) = _get_selected_report()

    document = Document()

    period = report['period']

    document.add_heading(
        'GRANDMASTER TECH ERP',
        level=0
    )

    document.add_heading(
        'Service & Repair Performance Report',
        level=1
    )

    document.add_paragraph(
        f"Reporting Period: {period['label']}"
    )

    document.add_heading(
        'Summary',
        level=1
    )

    summary = [
        ('Total Jobs', report['total_jobs']),
        ('Completed Jobs', report['completed_jobs']),
        ('Pending Jobs', report['pending_jobs']),
        ('In Progress Jobs', report['in_progress_jobs']),
        ('Completion Rate', f"{report['completion_rate']}%"),
        ('Customers', report['total_customers']),
        ('Equipment', report['total_equipment']),
    ]

    table = document.add_table(
        rows=1,
        cols=2
    )

    table.style = 'Table Grid'

    table.rows[0].cells[0].text = 'Metric'
    table.rows[0].cells[1].text = 'Value'

    for key, value in summary:
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = str(value)

    document.add_heading(
        'Financial Performance (MWK)',
        level=1
    )

    financial = [
        ('Service Charges', report['total_service_charges']),
        ('Parts Cost', report['total_parts_cost']),
        ('Total Revenue', report['total_revenue']),
        ('Amount Paid', report['total_amount_paid']),
        ('Outstanding Balance', report['total_balance']),
        ('Profit', report['total_profit']),
    ]

    table = document.add_table(
        rows=1,
        cols=2
    )

    table.style = 'Table Grid'

    table.rows[0].cells[0].text = 'Metric'
    table.rows[0].cells[1].text = 'Amount (MWK)'

    for key, value in financial:
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = f"{float(value or 0):,.2f}"

    document.add_heading(
        'Customer Demographics',
        level=1
    )

    table = document.add_table(
        rows=1,
        cols=3
    )

    table.style = 'Table Grid'

    table.rows[0].cells[0].text = 'Category'
    table.rows[0].cells[1].text = 'Value'
    table.rows[0].cells[2].text = 'Count'

    for category, data in [
        ('Gender', report['gender_counts']),
        ('Age Group', report['age_group_counts']),
        ('District', report['district_counts']),
        ('Customer Type', report['customer_type_counts']),
    ]:
        for key, value in data.items():
            cells = table.add_row().cells
            cells[0].text = category
            cells[1].text = str(key)
            cells[2].text = str(value)

    output = BytesIO()

    document.save(output)

    output.seek(0)

    filename = _report_filename(
        report_type,
        year,
        month,
        quarter,
        'docx'
    )

    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype=(
            'application/vnd.openxmlformats-officedocument.'
            'wordprocessingml.document'
        )
    )

'''

# ============================================================
# REPLACE OLD REPORT ROUTES WITH NEW ROUTES
# ============================================================

text = text[:start] + new_routes + text[end:]

app.write_text(text)

print("Part 4 complete.")
print("Excel export added.")
print("PDF export added.")
print("Word export added.")
print("All report periods supported.")

# ============================================================
# FINAL VALIDATION
# ============================================================

print("")
print("==============================================")
print("FULL REPORT EXPORT UPGRADE READY")
print("==============================================")
print("CSV   : /reports/download/csv")
print("Excel : /reports/download/excel")
print("PDF   : /reports/download/pdf")
print("Word  : /reports/download/word")
print("")
print("Supported report types:")
print("- Monthly")
print("- Quarterly")
print("- Yearly")
print("- All-Time")
print("")
print("Customer demographics included:")
print("- Gender")
print("- Age Group")
print("- District")
print("- Customer Type")
print("")
print("Financial analytics included:")
print("- Service Charges")
print("- Parts Cost")
print("- Total Revenue")
print("- Amount Paid")
print("- Outstanding Balance")
print("- Profit")
print("==============================================")

