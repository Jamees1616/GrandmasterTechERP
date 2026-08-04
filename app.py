from flask import Flask, render_template, request, redirect, url_for, flash, send_from_directory
from werkzeug.utils import secure_filename
import os
import requests
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'dev-only-secret-key')
database_url = os.environ.get('DATABASE_URL')
if database_url:
    if database_url.startswith('postgres://'):
        database_url = database_url.replace('postgres://', 'postgresql://', 1)
    app.config['SQLALCHEMY_DATABASE_URI'] = database_url
else:
    app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///grandmaster_tech.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = 'instance/uploads/job_attachments'
app.config['MAX_CONTENT_LENGTH'] = 25 * 1024 * 1024

# ========== SUPABASE STORAGE ==========
SUPABASE_URL = os.environ.get('SUPABASE_URL')
SUPABASE_SERVICE_KEY = os.environ.get('SUPABASE_SERVICE_KEY')
SUPABASE_BUCKET = 'job-attachments'


db = SQLAlchemy(app)


# ========== DATABASE MIGRATION ==========
def migrate_customer_demographics():
    """Add customer demographic columns if they do not exist."""
    from sqlalchemy import text

    try:
        with app.app_context():
            with db.engine.begin() as connection:
                columns = {
                    "gender": "VARCHAR(30)",
                    "age_group": "VARCHAR(30)",
                    "district": "VARCHAR(100)",
                    "customer_type": "VARCHAR(50)",
                }

                for column, column_type in columns.items():
                    connection.execute(
                        text(
                            f"ALTER TABLE customer "
                            f"ADD COLUMN IF NOT EXISTS {column} {column_type}"
                        )
                    )

                print("SUCCESS: Customer demographics database migration complete.")

    except Exception as e:
        print(f"WARNING: Customer demographics migration failed: {e}")


migrate_customer_demographics()

# ========== DATABASE MODELS ==========

class Customer(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    phone = db.Column(db.String(20))
    email = db.Column(db.String(100))
    company = db.Column(db.String(100))
    address = db.Column(db.Text)
    gender = db.Column(db.String(30))
    age_group = db.Column(db.String(30))
    district = db.Column(db.String(100))
    customer_type = db.Column(db.String(50))
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    equipment = db.relationship('Equipment', backref='customer', lazy=True)

class Equipment(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    customer_id = db.Column(db.Integer, db.ForeignKey('customer.id'), nullable=False)
    category = db.Column(db.String(50), nullable=False)
    brand = db.Column(db.String(50))
    model = db.Column(db.String(50))
    serial_number = db.Column(db.String(100))
    description = db.Column(db.Text)
    status = db.Column(db.String(20), default='Active')
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    job_cards = db.relationship('JobCard', backref='equipment', lazy=True)

class JobCard(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    equipment_id = db.Column(db.Integer, db.ForeignKey('equipment.id'), nullable=False)
    title = db.Column(db.String(200), nullable=False)
    problem_description = db.Column(db.Text)
    diagnosis = db.Column(db.Text)
    repair_process = db.Column(db.Text)
    parts_used = db.Column(db.Text)
    service_charge = db.Column(db.Numeric(12, 2), default=0)
    parts_cost = db.Column(db.Numeric(12, 2), default=0)
    total_cost = db.Column(db.Numeric(12, 2), default=0)
    amount_paid = db.Column(db.Numeric(12, 2), default=0)
    balance = db.Column(db.Numeric(12, 2), default=0)
    profit = db.Column(db.Numeric(12, 2), default=0)
    status = db.Column(db.String(20), default='Pending')
    priority = db.Column(db.String(20), default='Normal')
    started_at = db.Column(db.DateTime)
    completed_at = db.Column(db.DateTime)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

    attachments = db.relationship(
        'JobAttachment',
        backref='job_card',
        lazy=True,
        cascade='all, delete-orphan'
    )

class JobAttachment(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    job_card_id = db.Column(
        db.Integer,
        db.ForeignKey('job_card.id'),
        nullable=False
    )
    filename = db.Column(db.String(255), nullable=False)
    original_filename = db.Column(db.String(255))
    file_type = db.Column(db.String(100))
    category = db.Column(db.String(50), default='Other')
    uploaded_at = db.Column(
        db.DateTime,
        default=datetime.utcnow
    )


# ========== ROUTES ==========

@app.route('/')
def dashboard():
    total_customers = Customer.query.count()
    total_equipment = Equipment.query.count()
    total_jobs = JobCard.query.count()
    pending_jobs = JobCard.query.filter_by(status='Pending').count()
    in_progress_jobs = JobCard.query.filter_by(status='In Progress').count()
    completed_jobs = JobCard.query.filter_by(status='Completed').count()
    
    recent_jobs = JobCard.query.order_by(JobCard.created_at.desc()).limit(5).all()
    
    return render_template('dashboard.html', 
        total_customers=total_customers,
        total_equipment=total_equipment,
        total_jobs=total_jobs,
        pending_jobs=pending_jobs,
        in_progress_jobs=in_progress_jobs,
        completed_jobs=completed_jobs,
        recent_jobs=recent_jobs)

@app.route('/customers')
def customers():
    all_customers = Customer.query.order_by(Customer.created_at.desc()).all()
    return render_template('customers.html', customers=all_customers)

@app.route('/customers/<int:id>')
def customer_detail(id):
    customer = Customer.query.get_or_404(id)
    return render_template('customer_detail.html', customer=customer)


@app.route('/customers/<int:id>/update', methods=['POST'])
def update_customer(id):
    customer = Customer.query.get_or_404(id)

    customer.name = request.form.get('name', '').strip()
    customer.phone = request.form.get('phone', '').strip()
    customer.email = request.form.get('email', '').strip()
    customer.company = request.form.get('company', '').strip()
    customer.address = request.form.get('address', '').strip()
    customer.gender = request.form.get('gender', '').strip()
    customer.age_group = request.form.get('age_group', '').strip()
    customer.district = request.form.get('district', '').strip()
    customer.customer_type = request.form.get('customer_type', '').strip()

    if not customer.name:
        flash('Customer name is required.', 'danger')
        return redirect(url_for('customer_detail', id=id))

    db.session.commit()
    flash('Customer details updated successfully!', 'success')
    return redirect(url_for('customer_detail', id=id))

@app.route('/customers/add', methods=['POST'])
def add_customer():
    customer = Customer(
        name=request.form['name'],
        phone=request.form['phone'],
        email=request.form['email'],
        company=request.form['company'],
        address=request.form['address'],
        gender=request.form.get('gender', '').strip(),
        age_group=request.form.get('age_group', '').strip(),
        district=request.form.get('district', '').strip(),
        customer_type=request.form.get('customer_type', '').strip()
    )
    db.session.add(customer)
    db.session.commit()
    flash('Customer added successfully!', 'success')
    return redirect(url_for('customers'))

@app.route('/equipment')
def equipment():
    all_equipment = Equipment.query.order_by(Equipment.created_at.desc()).all()
    customers = Customer.query.all()
    return render_template('equipment.html', equipment=all_equipment, customers=customers)

@app.route('/equipment/add', methods=['POST'])
def add_equipment():
    equipment = Equipment(
        customer_id=request.form['customer_id'],
        category=request.form['category'],
        brand=request.form['brand'],
        model=request.form['model'],
        serial_number=request.form['serial_number'],
        description=request.form['description']
    )
    db.session.add(equipment)
    db.session.commit()
    flash('Equipment added successfully!', 'success')
    return redirect(url_for('equipment'))


@app.route('/equipment/<int:id>')
def equipment_detail(id):
    equipment = Equipment.query.get_or_404(id)
    return render_template('equipment_detail.html', equipment=equipment)

@app.route('/jobcards')
def jobcards():
    all_jobs = JobCard.query.order_by(JobCard.created_at.desc()).all()
    equipment_list = Equipment.query.all()
    return render_template('jobcards.html', jobs=all_jobs, equipment=equipment_list)

@app.route('/jobcards/add', methods=['POST'])
def add_jobcard():
    job = JobCard(
        equipment_id=request.form['equipment_id'],
        title=request.form['title'],
        problem_description=request.form['problem_description'],
        priority=request.form['priority'],
        status='Pending'
    )
    db.session.add(job)
    db.session.commit()
    flash('Job Card created successfully!', 'success')
    return redirect(url_for('jobcards'))

@app.route('/jobcards/update/<int:id>', methods=['POST'])
def update_jobcard(id):
    job = JobCard.query.get_or_404(id)

    old_status = job.status
    new_status = request.form.get('status', 'Pending')

    job.title = request.form.get('title', '')
    job.problem_description = request.form.get('problem_description', '')
    job.diagnosis = request.form.get('diagnosis', '')
    job.repair_process = request.form.get('repair_process', '')
    job.parts_used = request.form.get('parts_used', '')

    # Financial details (Malawian Kwacha)
    from decimal import Decimal, InvalidOperation

    def money_value(value):
        try:
            return Decimal(value or '0')
        except (InvalidOperation, ValueError):
            return Decimal('0')

    job.service_charge = money_value(request.form.get('service_charge'))
    job.parts_cost = money_value(request.form.get('parts_cost'))
    job.amount_paid = money_value(request.form.get('amount_paid'))

    # Automatic financial calculations
    job.total_cost = job.service_charge + job.parts_cost
    job.balance = job.total_cost - job.amount_paid
    job.profit = job.service_charge - job.parts_cost

    job.status = new_status
    job.priority = request.form.get('priority', 'Normal')

    # Record the moment work starts
    if new_status == 'In Progress' and not job.started_at:
        job.started_at = datetime.utcnow()

    # Record the moment the repair is completed
    if new_status == 'Completed' and not job.completed_at:
        job.completed_at = datetime.utcnow()

    # Reopen a completed repair
    if old_status == 'Completed' and new_status == 'In Progress':
        job.completed_at = None

    # Reset workflow timestamps if returned to Pending
    if new_status == 'Pending':
        job.started_at = None
        job.completed_at = None

    db.session.commit()

    flash('Job Card updated successfully!', 'success')
    return redirect(url_for('jobcards'))


@app.route('/jobcards/<int:id>/attachments/upload', methods=['POST'])
def upload_job_attachment(id):
    job = JobCard.query.get_or_404(id)

    uploaded_file = request.files.get('file')
    category = request.form.get('category', 'Other')

    if not uploaded_file or not uploaded_file.filename:
        flash('Please select a file to upload.', 'error')
        return redirect(url_for('jobcards'))

    original_filename = uploaded_file.filename
    safe_filename = secure_filename(original_filename)

    if not safe_filename:
        flash('Invalid file name.', 'error')
        return redirect(url_for('jobcards'))

    timestamp = datetime.utcnow().strftime('%Y%m%d%H%M%S')
    stored_filename = f"{job.id}_{timestamp}_{safe_filename}"

    # Upload file to Supabase Storage
    if not SUPABASE_URL or not SUPABASE_SERVICE_KEY:
        flash('Supabase Storage is not configured.', 'error')
        return redirect(url_for('jobcards'))

    supabase_upload_url = (
        f"{SUPABASE_URL}/storage/v1/object/"
        f"{SUPABASE_BUCKET}/{stored_filename}"
    )

    file_data = uploaded_file.read()

    supabase_response = requests.post(
        supabase_upload_url,
        headers={
            'Authorization': f'Bearer {SUPABASE_SERVICE_KEY}',
            'apikey': SUPABASE_SERVICE_KEY,
            'Content-Type': uploaded_file.mimetype or 'application/octet-stream',
        },
        data=file_data,
        timeout=60
    )

    if not supabase_response.ok:
        print("Supabase upload failed:", supabase_response.status_code, supabase_response.text)
        flash('Failed to upload attachment to storage.', 'error')
        return redirect(url_for('jobcards'))

    attachment = JobAttachment(
        job_card_id=job.id,
        filename=stored_filename,
        original_filename=original_filename,
        file_type=uploaded_file.mimetype,
        category=category
    )

    db.session.add(attachment)
    db.session.commit()

    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return {
            'success': True,
            'message': 'Attachment uploaded successfully!',
            'attachment': {
                'id': attachment.id,
                'filename': attachment.filename,
                'original_filename': attachment.original_filename,
                'category': attachment.category,
                'file_type': attachment.file_type,
                'view_url': url_for('view_job_attachment', filename=attachment.filename),
                'download_url': url_for('download_job_attachment', filename=attachment.filename)
            }
        }

    flash('Attachment uploaded successfully!', 'success')
    return redirect(url_for('jobcards'))


@app.route('/jobcards/attachments/<path:filename>')
def view_job_attachment(filename):
    if not SUPABASE_URL:
        return 'Supabase Storage is not configured.', 500

    file_url = (
        f"{SUPABASE_URL}/storage/v1/object/"
        f"{SUPABASE_BUCKET}/{filename}"
    )

    response = requests.get(
        file_url,
        headers={
            'Authorization': f'Bearer {SUPABASE_SERVICE_KEY}',
            'apikey': SUPABASE_SERVICE_KEY
        },
        timeout=60
    )

    if not response.ok:
        return 'Attachment not found.', 404

    from flask import Response

    return Response(
        response.content,
        status=200,
        content_type=response.headers.get(
            'Content-Type',
            'application/octet-stream'
        )
    )


@app.route('/jobcards/attachments/<path:filename>/download')
def download_job_attachment(filename):
    if not SUPABASE_URL:
        return 'Supabase Storage is not configured.', 500

    file_url = (
        f"{SUPABASE_URL}/storage/v1/object/"
        f"{SUPABASE_BUCKET}/{filename}"
    )

    response = requests.get(
        file_url,
        headers={
            'Authorization': f'Bearer {SUPABASE_SERVICE_KEY}',
            'apikey': SUPABASE_SERVICE_KEY
        },
        timeout=60
    )

    if not response.ok:
        return 'Attachment not found.', 404

    from flask import Response

    return Response(
        response.content,
        status=200,
        content_type=response.headers.get(
            'Content-Type',
            'application/octet-stream'
        ),
        headers={
            'Content-Disposition':
                f'attachment; filename="{filename}"'
        }
    )


@app.route('/jobcards/attachments/<int:attachment_id>/delete', methods=['POST'])
def delete_job_attachment(attachment_id):
    attachment = JobAttachment.query.get_or_404(attachment_id)

    # Delete file from Supabase Storage
    if SUPABASE_URL and SUPABASE_SERVICE_KEY:
        supabase_delete_url = (
            f"{SUPABASE_URL}/storage/v1/object/"
            f"{SUPABASE_BUCKET}/{attachment.filename}"
        )

        delete_response = requests.delete(
            supabase_delete_url,
            headers={
                'Authorization': f'Bearer {SUPABASE_SERVICE_KEY}',
                'apikey': SUPABASE_SERVICE_KEY
            },
            timeout=60
        )

        if not delete_response.ok:
            print(
                "Supabase delete failed:",
                delete_response.status_code,
                delete_response.text
            )

    db.session.delete(attachment)
    db.session.commit()

    flash('Attachment deleted successfully!', 'success')
    return redirect(url_for('jobcards'))



# ============================================================
# CENTRAL REPORTING ENGINE
# Supports: Monthly, Quarterly, Yearly, All-Time
# Currency: Malawian Kwacha (MWK / MK)
# ============================================================

def get_report_date_range(report_type='monthly', year=None, month=None, quarter=None):
    from datetime import datetime

    now = datetime.utcnow()

    year = year or now.year
    month = month or now.month

    report_type = (report_type or 'monthly').lower()

    if report_type == 'monthly':
        if month < 1 or month > 12:
            month = now.month

        start_date = datetime(year, month, 1)

        if month == 12:
            end_date = datetime(year + 1, 1, 1)
        else:
            end_date = datetime(year, month + 1, 1)

        label = f"{year}-{month:02d}"

    elif report_type == 'quarterly':
        quarter = quarter or ((month - 1) // 3 + 1)

        if quarter < 1 or quarter > 4:
            quarter = 1

        start_month = ((quarter - 1) * 3) + 1
        start_date = datetime(year, start_month, 1)

        if quarter == 4:
            end_date = datetime(year + 1, 1, 1)
        else:
            end_date = datetime(year, start_month + 3, 1)

        label = f"Q{quarter} {year}"

    elif report_type == 'yearly':
        start_date = datetime(year, 1, 1)
        end_date = datetime(year + 1, 1, 1)

        label = str(year)

    elif report_type == 'all':
        start_date = None
        end_date = None

        label = "All Time"

    else:
        raise ValueError(f"Unsupported report type: {report_type}")

    return {
        'type': report_type,
        'year': year,
        'month': month,
        'quarter': quarter,
        'start_date': start_date,
        'end_date': end_date,
        'label': label
    }


def get_report_data(report_type='monthly', year=None, month=None, quarter=None):
    period = get_report_date_range(
        report_type=report_type,
        year=year,
        month=month,
        quarter=quarter
    )

    job_query = JobCard.query
    customer_query = Customer.query
    equipment_query = Equipment.query

    if period['start_date'] and period['end_date']:
        job_query = job_query.filter(
            JobCard.created_at >= period['start_date'],
            JobCard.created_at < period['end_date']
        )

        customer_query = customer_query.filter(
            Customer.created_at >= period['start_date'],
            Customer.created_at < period['end_date']
        )

        equipment_query = equipment_query.filter(
            Equipment.created_at >= period['start_date'],
            Equipment.created_at < period['end_date']
        )

    jobs = job_query.all()
    customers = customer_query.all()
    equipment = equipment_query.all()

    total_jobs = len(jobs)

    pending_jobs = sum(
        1 for job in jobs if job.status == 'Pending'
    )

    in_progress_jobs = sum(
        1 for job in jobs if job.status == 'In Progress'
    )

    completed_jobs = sum(
        1 for job in jobs if job.status == 'Completed'
    )

    completion_rate = (
        round((completed_jobs / total_jobs) * 100, 1)
        if total_jobs else 0
    )

    total_service_charges = sum(
        (job.service_charge or 0) for job in jobs
    )

    total_parts_cost = sum(
        (job.parts_cost or 0) for job in jobs
    )

    total_revenue = sum(
        (job.total_cost or 0) for job in jobs
    )

    total_amount_paid = sum(
        (job.amount_paid or 0) for job in jobs
    )

    total_balance = sum(
        (job.balance or 0) for job in jobs
    )

    total_profit = sum(
        (job.profit or 0) for job in jobs
    )

    priority_counts = {}

    for job in jobs:
        priority = job.priority or 'Normal'

        priority_counts[priority] = (
            priority_counts.get(priority, 0) + 1
        )

    category_counts = {}

    for job in jobs:
        category = (
            job.equipment.category
            if job.equipment
            else 'Unknown'
        )

        category_counts[category] = (
            category_counts.get(category, 0) + 1
        )

    customer_counts = {}

    for job in jobs:
        if job.equipment and job.equipment.customer:
            customer_name = job.equipment.customer.name

            customer_counts[customer_name] = (
                customer_counts.get(customer_name, 0) + 1
            )

    top_customers = sorted(
        customer_counts.items(),
        key=lambda item: item[1],
        reverse=True
    )[:10]

    # Customer demographic analytics
    gender_counts = {}
    age_group_counts = {}
    district_counts = {}
    customer_type_counts = {}

    for customer in customers:
        gender = customer.gender or 'Unknown'
        age_group = customer.age_group or 'Unknown'
        district = customer.district or 'Unknown'
        customer_type = customer.customer_type or 'Unknown'

        gender_counts[gender] = gender_counts.get(gender, 0) + 1
        age_group_counts[age_group] = age_group_counts.get(age_group, 0) + 1
        district_counts[district] = district_counts.get(district, 0) + 1
        customer_type_counts[customer_type] = (
            customer_type_counts.get(customer_type, 0) + 1
        )

    return {
        'period': period,
        'jobs': jobs,
        'customers': customers,
        'equipment': equipment,

        'total_jobs': total_jobs,
        'total_customers': len(customers),
        'total_equipment': len(equipment),

        'pending_jobs': pending_jobs,
        'in_progress_jobs': in_progress_jobs,
        'completed_jobs': completed_jobs,
        'completion_rate': completion_rate,

        # Financial analytics — MWK
        'total_service_charges': total_service_charges,
        'total_parts_cost': total_parts_cost,
        'total_revenue': total_revenue,
        'total_amount_paid': total_amount_paid,
        'total_balance': total_balance,
        'total_profit': total_profit,

        'priority_counts': priority_counts,
        'category_counts': category_counts,
        'top_customers': top_customers,

        # Customer demographic analytics
        'gender_counts': gender_counts,
        'age_group_counts': age_group_counts,
        'district_counts': district_counts,
        'customer_type_counts': customer_type_counts
    }




# ============================================================
# REPORT DOWNLOAD / EXPORT ENGINE
# Supports Monthly, Quarterly, Yearly, All-Time
# Formats: CSV, Excel, PDF, Word
# Currency: Malawian Kwacha (MWK / MK)
# ============================================================

def _report_args():
    from datetime import datetime

    now = datetime.utcnow()

    report_type = request.args.get(
        'report_type',
        'monthly'
    ).lower()

    year = request.args.get(
        'year',
        now.year,
        type=int
    )

    month = request.args.get(
        'month',
        now.month,
        type=int
    )

    quarter = request.args.get(
        'quarter',
        1,
        type=int
    )

    if report_type not in [
        'monthly',
        'quarterly',
        'yearly',
        'all'
    ]:
        report_type = 'monthly'

    if month < 1 or month > 12:
        month = now.month

    if quarter < 1 or quarter > 4:
        quarter = 1

    return (
        report_type,
        year,
        month,
        quarter
    )


def _report_filename(
    report_type,
    year,
    month,
    quarter,
    extension
):
    period = get_report_date_range(
        report_type=report_type,
        year=year,
        month=month,
        quarter=quarter
    )

    label = (
        period['label']
        .replace(' ', '_')
        .replace('/', '-')
    )

    return (
        f"Grandmaster_Tech_ERP_Report_"
        f"{label}.{extension}"
    )


def _get_selected_report():
    report_type, year, month, quarter = _report_args()

    report = get_report_data(
        report_type=report_type,
        year=year,
        month=month,
        quarter=quarter
    )

    return (
        report_type,
        year,
        month,
        quarter,
        report
    )


@app.route('/reports')
def reports():
    (
        report_type,
        selected_year,
        selected_month,
        selected_quarter,
        report
    ) = _get_selected_report()

    return render_template(
        'reports.html',

        report_type=report_type,
        selected_month=selected_month,
        selected_year=selected_year,
        selected_quarter=selected_quarter,

        report_period_label=report['period']['label'],

        total_jobs=report['total_jobs'],
        total_customers=report['total_customers'],
        total_equipment=report['total_equipment'],

        pending_jobs=report['pending_jobs'],
        in_progress_jobs=report['in_progress_jobs'],
        completed_jobs=report['completed_jobs'],
        completion_rate=report['completion_rate'],

        total_service_charges=report['total_service_charges'],
        total_parts_cost=report['total_parts_cost'],
        total_revenue=report['total_revenue'],
        total_amount_paid=report['total_amount_paid'],
        total_balance=report['total_balance'],
        total_profit=report['total_profit'],

        priority_counts=report['priority_counts'],
        category_counts=report['category_counts'],
        top_customers=report['top_customers'],

        gender_counts=report['gender_counts'],
        age_group_counts=report['age_group_counts'],
        district_counts=report['district_counts'],
        customer_type_counts=report['customer_type_counts']
    )


# ============================================================
# CSV EXPORT
# ============================================================

@app.route('/reports/download/csv')
def download_reports_csv():
    import csv
    from io import StringIO
    from flask import Response

    (
        report_type,
        year,
        month,
        quarter,
        report
    ) = _get_selected_report()

    output = StringIO()
    writer = csv.writer(output)

    period = report['period']

    writer.writerow([
        'GRANDMASTER TECH ERP'
    ])

    writer.writerow([
        'Service & Repair Performance Report'
    ])

    writer.writerow([
        'Report Period',
        period['label']
    ])

    writer.writerow([
        'Report Type',
        period['type'].title()
    ])

    writer.writerow([])

    writer.writerow([
        'SUMMARY'
    ])

    writer.writerow([
        'Total Jobs',
        report['total_jobs']
    ])

    writer.writerow([
        'Completed Jobs',
        report['completed_jobs']
    ])

    writer.writerow([
        'Pending Jobs',
        report['pending_jobs']
    ])

    writer.writerow([
        'In Progress Jobs',
        report['in_progress_jobs']
    ])

    writer.writerow([
        'Completion Rate',
        f"{report['completion_rate']}%"
    ])

    writer.writerow([
        'Customers',
        report['total_customers']
    ])

    writer.writerow([
        'Equipment',
        report['total_equipment']
    ])

    writer.writerow([])

    writer.writerow([
        'FINANCIAL PERFORMANCE - MWK'
    ])

    writer.writerow([
        'Service Charges',
        report['total_service_charges']
    ])

    writer.writerow([
        'Parts Cost',
        report['total_parts_cost']
    ])

    writer.writerow([
        'Total Revenue',
        report['total_revenue']
    ])

    writer.writerow([
        'Amount Paid',
        report['total_amount_paid']
    ])

    writer.writerow([
        'Outstanding Balance',
        report['total_balance']
    ])

    writer.writerow([
        'Profit',
        report['total_profit']
    ])

    writer.writerow([])

    writer.writerow([
        'CUSTOMER DEMOGRAPHICS'
    ])

    writer.writerow([
        'Gender',
        'Count'
    ])

    for key, value in report[
        'gender_counts'
    ].items():
        writer.writerow([
            key,
            value
        ])

    writer.writerow([])

    writer.writerow([
        'Age Group',
        'Count'
    ])

    for key, value in report[
        'age_group_counts'
    ].items():
        writer.writerow([
            key,
            value
        ])

    writer.writerow([])

    writer.writerow([
        'District',
        'Count'
    ])

    for key, value in report[
        'district_counts'
    ].items():
        writer.writerow([
            key,
            value
        ])

    writer.writerow([])

    writer.writerow([
        'Customer Type',
        'Count'
    ])

    for key, value in report[
        'customer_type_counts'
    ].items():
        writer.writerow([
            key,
            value
        ])

    writer.writerow([])

    writer.writerow([
        'JOB DETAILS'
    ])

    writer.writerow([
        'ID',
        'Job Title',
        'Equipment Category',
        'Brand',
        'Model',
        'Status',
        'Priority',
        'Created Date',
        'Service Charge (MWK)',
        'Parts Cost (MWK)',
        'Total Cost (MWK)',
        'Amount Paid (MWK)',
        'Balance (MWK)',
        'Profit (MWK)'
    ])

    for job in report['jobs']:

        equipment = job.equipment

        writer.writerow([
            job.id,
            job.title,
            equipment.category
            if equipment else '',
            equipment.brand
            if equipment else '',
            equipment.model
            if equipment else '',
            job.status,
            job.priority,
            job.created_at.strftime(
                '%Y-%m-%d'
            )
            if job.created_at
            else '',
            job.service_charge or 0,
            job.parts_cost or 0,
            job.total_cost or 0,
            job.amount_paid or 0,
            job.balance or 0,
            job.profit or 0
        ])

    filename = _report_filename(
        report_type,
        year,
        month,
        quarter,
        'csv'
    )

    return Response(
        output.getvalue(),
        mimetype='text/csv',
        headers={
            'Content-Disposition':
            f'attachment; filename="{filename}"'
        }
    )


@app.route('/reports/download/excel')
def download_reports_excel():
    from io import BytesIO
    from flask import send_file
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment

    (
        report_type,
        year,
        month,
        quarter,
        report
    ) = _get_selected_report()

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Report"

    period = report['period']

    rows = [
        ['GRANDMASTER TECH ERP'],
        ['Service & Repair Performance Report'],
        ['Report Period', period['label']],
        ['Report Type', period['type'].title()],
        [],
        ['SUMMARY'],
        ['Total Jobs', report['total_jobs']],
        ['Completed Jobs', report['completed_jobs']],
        ['Pending Jobs', report['pending_jobs']],
        ['In Progress Jobs', report['in_progress_jobs']],
        ['Completion Rate', f"{report['completion_rate']}%"],
        ['Customers', report['total_customers']],
        ['Equipment', report['total_equipment']],
        [],
        ['FINANCIAL PERFORMANCE - MWK'],
        ['Service Charges', report['total_service_charges']],
        ['Parts Cost', report['total_parts_cost']],
        ['Total Revenue', report['total_revenue']],
        ['Amount Paid', report['total_amount_paid']],
        ['Outstanding Balance', report['total_balance']],
        ['Profit', report['total_profit']],
        [],
        ['CUSTOMER DEMOGRAPHICS'],
        ['Gender', 'Count']
    ]

    for key, value in report['gender_counts'].items():
        rows.append([key, value])

    rows += [
        [],
        ['Age Group', 'Count']
    ]

    for key, value in report['age_group_counts'].items():
        rows.append([key, value])

    rows += [
        [],
        ['District', 'Count']
    ]

    for key, value in report['district_counts'].items():
        rows.append([key, value])

    rows += [
        [],
        ['Customer Type', 'Count']
    ]

    for key, value in report['customer_type_counts'].items():
        rows.append([key, value])

    rows += [
        [],
        ['JOB DETAILS'],
        [
            'ID',
            'Job Title',
            'Equipment Category',
            'Brand',
            'Model',
            'Status',
            'Priority',
            'Created Date',
            'Service Charge (MWK)',
            'Parts Cost (MWK)',
            'Total Cost (MWK)',
            'Amount Paid (MWK)',
            'Balance (MWK)',
            'Profit (MWK)'
        ]
    ]

    for job in report['jobs']:

        equipment = job.equipment

        rows.append([
            job.id,
            job.title,
            equipment.category if equipment else '',
            equipment.brand if equipment else '',
            equipment.model if equipment else '',
            job.status,
            job.priority,
            job.created_at.strftime('%Y-%m-%d')
            if job.created_at else '',
            float(job.service_charge or 0),
            float(job.parts_cost or 0),
            float(job.total_cost or 0),
            float(job.amount_paid or 0),
            float(job.balance or 0),
            float(job.profit or 0)
        ])

    for row in rows:
        sheet.append(row)

    for cell in sheet[1]:
        cell.font = Font(
            bold=True,
            size=16
        )

    for cell in sheet[2]:
        cell.font = Font(
            bold=True
        )

    for row in sheet.iter_rows():
        for cell in row:
            cell.alignment = Alignment(
                vertical='top'
            )

    for column in sheet.columns:
        max_length = 0
        column_letter = column[0].column_letter

        for cell in column:
            if cell.value is not None:
                max_length = max(
                    max_length,
                    len(str(cell.value))
                )

        sheet.column_dimensions[
            column_letter
        ].width = min(
            max(max_length + 2, 12),
            40
        )

    output = BytesIO()

    workbook.save(output)

    output.seek(0)

    filename = _report_filename(
        report_type,
        year,
        month,
        quarter,
        'xlsx'
    )

    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype=(
            'application/vnd.openxmlformats-officedocument.'
            'spreadsheetml.sheet'
        )
    )


# ============================================================
# PDF EXPORT
# ============================================================

@app.route('/reports/download/pdf')
def download_reports_pdf():
    from io import BytesIO
    from flask import send_file
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle
    )
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet

    (
        report_type,
        year,
        month,
        quarter,
        report
    ) = _get_selected_report()

    output = BytesIO()

    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=36,
        leftMargin=36,
        topMargin=36,
        bottomMargin=36
    )

    styles = getSampleStyleSheet()

    story = []

    period = report['period']

    story.append(
        Paragraph(
            'GRANDMASTER TECH ERP',
            styles['Title']
        )
    )

    story.append(
        Paragraph(
            'Service & Repair Performance Report',
            styles['Heading2']
        )
    )

    story.append(
        Paragraph(
            f"Reporting Period: {period['label']}",
            styles['Normal']
        )
    )

    story.append(Spacer(1, 18))

    summary_data = [
        ['Metric', 'Value'],
        ['Total Jobs', report['total_jobs']],
        ['Completed Jobs', report['completed_jobs']],
        ['Pending Jobs', report['pending_jobs']],
        ['In Progress Jobs', report['in_progress_jobs']],
        ['Completion Rate', f"{report['completion_rate']}%"],
        ['Customers', report['total_customers']],
        ['Equipment', report['total_equipment']],
    ]

    table = Table(
        summary_data,
        colWidths=[250, 150]
    )

    table.setStyle(
        TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('PADDING', (0, 0), (-1, -1), 6),
        ])
    )

    story.append(table)

    story.append(Spacer(1, 18))

    story.append(
        Paragraph(
            'Financial Performance (MWK)',
            styles['Heading2']
        )
    )

    financial_data = [
        ['Metric', 'Amount (MWK)'],
        ['Service Charges', f"{report['total_service_charges']:,.2f}"],
        ['Parts Cost', f"{report['total_parts_cost']:,.2f}"],
        ['Total Revenue', f"{report['total_revenue']:,.2f}"],
        ['Amount Paid', f"{report['total_amount_paid']:,.2f}"],
        ['Outstanding Balance', f"{report['total_balance']:,.2f}"],
        ['Profit', f"{report['total_profit']:,.2f}"],
    ]

    financial_table = Table(
        financial_data,
        colWidths=[250, 150]
    )

    financial_table.setStyle(
        TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('PADDING', (0, 0), (-1, -1), 6),
        ])
    )

    story.append(financial_table)

    story.append(Spacer(1, 18))

    story.append(
        Paragraph(
            'Customer Demographics',
            styles['Heading2']
        )
    )

    demographic_data = [
        ['Category', 'Value', 'Count']
    ]

    for key, value in report[
        'gender_counts'
    ].items():
        demographic_data.append([
            'Gender',
            key,
            value
        ])

    for key, value in report[
        'age_group_counts'
    ].items():
        demographic_data.append([
            'Age Group',
            key,
            value
        ])

    for key, value in report[
        'district_counts'
    ].items():
        demographic_data.append([
            'District',
            key,
            value
        ])

    for key, value in report[
        'customer_type_counts'
    ].items():
        demographic_data.append([
            'Customer Type',
            key,
            value
        ])

    demographic_table = Table(
        demographic_data,
        colWidths=[120, 180, 100]
    )

    demographic_table.setStyle(
        TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('PADDING', (0, 0), (-1, -1), 5),
        ])
    )

    story.append(demographic_table)

    document.build(story)

    output.seek(0)

    filename = _report_filename(
        report_type,
        year,
        month,
        quarter,
        'pdf'
    )

    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype='application/pdf'
    )


# ============================================================
# WORD EXPORT
# ============================================================

@app.route('/reports/download/word')
def download_reports_word():
    from io import BytesIO
    from flask import send_file
    from docx import Document

    (
        report_type,
        year,
        month,
        quarter,
        report
    ) = _get_selected_report()

    document = Document()

    period = report['period']

    document.add_heading(
        'GRANDMASTER TECH ERP',
        level=0
    )

    document.add_heading(
        'Service & Repair Performance Report',
        level=1
    )

    document.add_paragraph(
        f"Reporting Period: {period['label']}"
    )

    document.add_heading(
        'Summary',
        level=1
    )

    summary = [
        ('Total Jobs', report['total_jobs']),
        ('Completed Jobs', report['completed_jobs']),
        ('Pending Jobs', report['pending_jobs']),
        ('In Progress Jobs', report['in_progress_jobs']),
        ('Completion Rate', f"{report['completion_rate']}%"),
        ('Customers', report['total_customers']),
        ('Equipment', report['total_equipment']),
    ]

    table = document.add_table(
        rows=1,
        cols=2
    )

    table.style = 'Table Grid'

    table.rows[0].cells[0].text = 'Metric'
    table.rows[0].cells[1].text = 'Value'

    for key, value in summary:
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = str(value)

    document.add_heading(
        'Financial Performance (MWK)',
        level=1
    )

    financial = [
        ('Service Charges', report['total_service_charges']),
        ('Parts Cost', report['total_parts_cost']),
        ('Total Revenue', report['total_revenue']),
        ('Amount Paid', report['total_amount_paid']),
        ('Outstanding Balance', report['total_balance']),
        ('Profit', report['total_profit']),
    ]

    table = document.add_table(
        rows=1,
        cols=2
    )

    table.style = 'Table Grid'

    table.rows[0].cells[0].text = 'Metric'
    table.rows[0].cells[1].text = 'Amount (MWK)'

    for key, value in financial:
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = f"{float(value or 0):,.2f}"

    document.add_heading(
        'Customer Demographics',
        level=1
    )

    table = document.add_table(
        rows=1,
        cols=3
    )

    table.style = 'Table Grid'

    table.rows[0].cells[0].text = 'Category'
    table.rows[0].cells[1].text = 'Value'
    table.rows[0].cells[2].text = 'Count'

    for category, data in [
        ('Gender', report['gender_counts']),
        ('Age Group', report['age_group_counts']),
        ('District', report['district_counts']),
        ('Customer Type', report['customer_type_counts']),
    ]:
        for key, value in data.items():
            cells = table.add_row().cells
            cells[0].text = category
            cells[1].text = str(key)
            cells[2].text = str(value)

    output = BytesIO()

    document.save(output)

    output.seek(0)

    filename = _report_filename(
        report_type,
        year,
        month,
        quarter,
        'docx'
    )

    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype=(
            'application/vnd.openxmlformats-officedocument.'
            'wordprocessingml.document'
        )
    )

@app.route('/knowledge')
def knowledge():
    return render_template('knowledge.html')

# ========== INIT ==========

def initialize_database():
    with app.app_context():
        db.create_all()

initialize_database()

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
